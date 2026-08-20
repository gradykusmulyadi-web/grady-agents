"""Build the Malaysia Market Entry project charter (2-page .docx).

Regenerable: edit the CONTENT constants below and re-run. Writes to the path given
as argv[1] (default: ./Malaysia_Expansion_Project_Charter_v1.docx).
"""
import sys

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ------------------------------------------------------------------ style kit
FONT = "Arial"
NAVY = "1F3864"
LIGHT = "D9E2F3"
FAINT = "F2F2F2"
RULE = "BFBFBF"

C_NAVY = RGBColor(0x1F, 0x38, 0x64)
C_GREY = RGBColor(0x59, 0x59, 0x59)
C_BLACK = RGBColor(0x00, 0x00, 0x00)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

SZ_TITLE = 16
SZ_SUB = 8.5
SZ_HEAD = 10.5
SZ_BODY = 9.5
SZ_TBL = 8.5
SZ_NOTE = 7.5

USABLE_CM = 18.6  # A4 (21.0cm) minus 2 x 1.2cm margins


# ------------------------------------------------------------------ xml helpers
def shade(cell, hexcolor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(shd)


def cell_margins(table, top=28, bottom=28, left=68, right=68):
    """Tighten table cell padding. Units are twentieths of a point."""
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    table._tbl.tblPr.append(mar)


def borders(table, color=RULE, sz=4, inner=True):
    el = OxmlElement("w:tblBorders")
    sides = ["top", "left", "bottom", "right"]
    if inner:
        sides += ["insideH", "insideV"]
    for side in sides:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:color"), color)
        el.append(b)
    table._tbl.tblPr.append(el)


def no_borders(table):
    el = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        el.append(b)
    table._tbl.tblPr.append(el)


def keep_together(par):
    """Stop Word orphaning a heading at the foot of a page."""
    pPr = par._p.get_or_add_pPr()
    for tag in ("w:keepNext", "w:keepLines"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), "true")
        pPr.append(el)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


# ------------------------------------------------------------------ writers
def write(par, text, *, size=SZ_BODY, bold=False, italic=False, color=C_BLACK):
    run = par.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run


def para(doc, text="", *, size=SZ_BODY, bold=False, italic=False, color=C_BLACK,
         before=0, after=3, align=None, line=1.0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        pf.alignment = align
    if text:
        write(p, text, size=size, bold=bold, italic=italic, color=color)
    return p


def section(doc, number, title, *, before=7):
    """Full-width navy bar heading. A shaded *paragraph*, not a table — two adjacent
    tables get merged into one by Word, and most sections lead straight into a table."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.0
    pf.left_indent = Cm(0.12)
    pf.right_indent = Cm(0.12)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), NAVY)
    pPr.append(shd)
    write(p, f"{number}   {title}", size=SZ_HEAD, bold=True, color=C_WHITE)
    keep_together(p)
    return p


def bullets(doc, items, *, size=SZ_BODY, indent=0.45):
    for it in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(1.5)
        pf.line_spacing = 1.0
        pf.left_indent = Cm(indent + 0.3)
        pf.first_line_indent = Cm(-0.3)
        write(p, "•  ", size=size)
        # bold text wrapped in **...**
        for i, chunk in enumerate(it.split("**")):
            if chunk:
                write(p, chunk, size=size, bold=(i % 2 == 1))


def table(doc, rows, widths_cm, *, header=True, size=SZ_TBL, align_right=None,
          bold_cols=(), zebra=False, after=0):
    """rows[0] is the header. Cell text supports **bold** spans."""
    align_right = align_right or set()
    t = doc.add_table(rows=len(rows), cols=len(widths_cm))
    t.autofit = False
    borders(t)
    cell_margins(t)
    for r, row in enumerate(rows):
        tr = t.rows[r]
        if header and r == 0:
            repeat_header(tr)
        for c, text in enumerate(row):
            cell = t.cell(r, c)
            cell.width = Cm(widths_cm[c])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if header and r == 0:
                shade(cell, NAVY)
            elif zebra and r % 2 == 0:
                shade(cell, FAINT)
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.0
            if c in align_right:
                pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hdr = header and r == 0
            base_bold = hdr or (c in bold_cols)
            color = C_WHITE if hdr else C_BLACK
            for i, chunk in enumerate(str(text).split("**")):
                if chunk:
                    write(p, chunk, size=size, bold=base_bold or (i % 2 == 1),
                          color=color)
    # column widths must also be declared on the grid
    grid = t._tbl.find(qn("w:tblGrid"))
    for col, w in zip(grid.findall(qn("w:gridCol")), widths_cm):
        col.set(qn("w:w"), str(int(w * 567)))
    if after:
        para(doc, "", size=after, after=0)
    return t


# ------------------------------------------------------------------ document
def build(path):
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(SZ_BODY)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0

    s = doc.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.top_margin = s.bottom_margin = Cm(1.1)
    s.left_margin = s.right_margin = Cm(1.2)

    # ---------------------------------------------------------- title block
    p = para(doc, after=0)
    write(p, "PROJECT CHARTER", size=SZ_TITLE, bold=True, color=C_NAVY)
    p = para(doc, after=2)
    write(p, "Malaysia Market Entry — SEA Expansion Beachhead",
          size=11, bold=True, color=C_BLACK)
    para(doc, "McEasy · Charter horizon Q3 2026 – Q4 2027 · All figures in USD "
              "unless stated otherwise",
         size=SZ_SUB, italic=True, color=C_GREY, after=4)

    meta = table(doc, [
        ["Sponsor", "Chief Operating Officer", "Project Owner",
         "Chief Strategy Officer"],
        ["Version", "v1.0 — Draft for review", "Date",
         "17 August 2026"],
        ["Status", "For COO / C-Level review", "Next gate",
         "Gate 1 — 31 December 2026"],
    ], [2.4, 6.9, 2.4, 6.9], header=False, size=SZ_TBL, bold_cols=(0, 2), zebra=False)
    for r in range(3):
        for c in (0, 2):
            shade(meta.cell(r, c), LIGHT)

    # ---------------------------------------------------------- 1. purpose
    section(doc, "1", "PURPOSE & BUSINESS CASE", before=6)
    para(doc, "Malaysia is the beachhead for McEasy's SEA & ANZ expansion — the first of "
              "eight market waves building toward USD 16.5m ARR by 2030. It is selected for the "
              "highest ARPU in near-ASEAN, a shared regulatory idiom, and an installed base of "
              "Indonesian logistics multinationals already operating there, which permits "
              "land-and-expand on existing relationships rather than cold entry.",
         before=3, after=2)
    p = para(doc, after=0)
    write(p, "Demand is regulatory, not discretionary. ", bold=True)
    write(p, "Under the APAD / JPJ Industry Code of Practice, commercial vehicles of 7.5 tonnes "
             "and above, customs-transit vehicles, tobacco and alcohol haulage, and all bus fleet "
             "operators must run a functioning GPS tracker to obtain or renew their operating "
             "permit. McEasy sells into an existing compliance obligation, not a new budget line.")

    # ---------------------------------------------------------- 2. objective
    section(doc, "2", "OBJECTIVE & SUCCESS CRITERIA")
    para(doc, "Deliver USD 500,000 cumulative ARR in Malaysia by 31 December 2027 on a repeatable "
              "commercial engine that materially reduces the cost and time of the next SEA market "
              "entry.", before=3, after=3)
    table(doc, [
        ["Success measure", "Q2 2027", "Q3 2027", "Q4 2027"],
        ["Cumulative ARR (USD)", "100,000", "275,000", "**500,000**"],
        ["Implied active units", "~595", "~1,637", "**~2,976**"],
        ["Implied customers", "~5", "~13", "**~24**"],
    ], [7.2, 3.8, 3.8, 3.8], align_right={1, 2, 3}, bold_cols=(0,), zebra=True)
    para(doc, "Derived at USD 14.00/month blended ARPU. Target mix is 50:50 large-fleet to "
              "mid-market — approximately 3–4 anchor logos of ~400 vehicles and ~20 "
              "mid-market accounts of ~75 vehicles, seeded from Indonesian customers with "
              "Malaysian operations.",
         size=SZ_NOTE, italic=True, color=C_GREY, before=2, after=0)

    # ---------------------------------------------------------- 3. scope
    section(doc, "3", "SCOPE")
    table(doc, [
        ["**In scope**",
         "Malaysia only. Entity and legal setup; device SKU selection and certification; MEP "
         "localisation (BM/EN, MYR, LHDN e-invoicing, Malaysian mapping); third-party installer "
         "network; sales, CS and implementation hiring; pricing and playbook documentation."],
        ["**Out of scope**",
         "Vietnam, Thailand, Philippines, Singapore and Australia entry (post-Q4 2027); "
         "acquisition of a Malaysian telematics provider (separate evaluation — see §5); "
         "hardware NPD beyond SKU selection and certification of existing devices; roadmap items "
         "not required for Malaysia fit."],
    ], [2.6, 16.0], header=False)
    for r in range(2):
        shade(doc.tables[-1].cell(r, 0), LIGHT)

    # ---------------------------------------------------------- 4. regulatory
    section(doc, "4", "REGULATORY & CERTIFICATION PATH")
    table(doc, [
        ["Layer", "Status", "Implication for the plan"],
        ["MCMC / SIRIM QAS type approval", "**Mandatory**",
         "Hard legal gate on importing and selling the device. Critical path — file Q4 2026, "
         "not Q1 2027."],
        ["APAD / JPJ ICOP conformance", "**Mandatory for the customer**",
         "The demand driver. Device must meet JPJ/APAD technical specification to carry permit "
         "value for the operator."],
        ["MIROS TrackScore star rating", "**Voluntary**",
         "Commercial table stakes — incumbents market 5-star ratings. Target ≥4-star by "
         "Q1 2027 to qualify in tenders."],
        ["MOT telematics initiative", "**Not yet mandatory**",
         "2026–27 advocacy → 2027–28 driver database → 2028+ possible mandate. "
         "Five providers are already in the pilot cohort."],
    ], [5.0, 3.6, 10.0], zebra=True)
    para(doc, "Sources are secondary (industry and vendor publications) and are treated as "
              "indicative. The Q3 2026 research milestone replaces them with primary confirmation "
              "from SIRIM, APAD and MIROS. The 500k target does not depend on the 2028 mandate "
              "arriving.",
         size=SZ_NOTE, italic=True, color=C_GREY, before=2, after=0)

    # ---------------------------------------------------------- 5. approach
    section(doc, "5", "APPROACH & GATE 1 DECISION (Q4 2026)")
    para(doc, "The recommended default is organic entry: a Malaysian partner-of-record for entity "
              "and market access, McEasy-employed sales and customer success, and a contracted "
              "third-party installer network. Acquiring a local player is excluded from this gate "
              "— diligence, SPA and regulatory approval cannot complete in time to preserve a "
              "Q1 2027 launch, and it should be evaluated as a separate initiative.",
         before=3, after=2)
    para(doc, "Gate 1 releases the launch only if all of the following hold by 31 December 2026:",
         bold=True, after=2)
    bullets(doc, [
        "Named pipeline of **≥5 Indonesian logos with Malaysian operations**, ≥800 "
        "vehicles combined, confirming intent",
        "**SIRIM / MCMC type-approval path, cost and lead time confirmed in writing**; SKU "
        "gap-assessed against a TrackScore ≥4-star target",
        "MOT pilot participation tested — open or closed — with competitive implications "
        "documented",
        "Incumbent pricing validated against the USD 14/month assumption",
        "≥2 qualified country-lead candidates identified",
    ])
    p = para(doc, before=2, after=0)
    write(p, "If unmet: ", bold=True)
    write(p, "defer launch by one quarter. Do not enter blind.")

    # ======================================================== PAGE 2
    br = doc.add_paragraph()
    br.paragraph_format.space_after = Pt(0)
    br.add_run().add_break(WD_BREAK.PAGE)

    section(doc, "6", "MILESTONES & EXIT CRITERIA", before=0)
    table(doc, [
        ["Quarter", "Milestone", "Exit criteria"],
        ["Q3 2026", "Market validation",
         "≥8 practitioner interviews (Malaysian operators, ex-telematics); ≥5 named seed "
         "accounts; pricing validated; regulatory stack confirmed (SIRIM/MCMC, APAD ICOP spec, "
         "TrackScore criteria, MOT pilot status); TH/PH/VN diligence logged for later waves"],
        ["Q4 2026", "**Gate 1 — GTM mode**",
         "Entry mode signed by sponsor; entity setup initiated; device SKU frozen by November; "
         "**SIRIM/MCMC filed and TrackScore assessment submitted**"],
        ["Q1 2027", "Market-ready",
         "Type approval granted; ≥4-star TrackScore achieved; SKU conforms to APAD/JPJ ICOP; "
         "MEP localisation live; **≥4 installer partners contracted**; country lead and "
         "enterprise AE onboarded"],
        ["Q2 2027", "**Gate 2 — First revenue**",
         "First contract signed; **100k cumulative ARR**; ≥1 anchor logo live; installer SPV "
         "and SLA operational"],
        ["Q3 2027", "Scale",
         "**275k cumulative ARR**; 2 mid-market AEs productive; product representative in-country; "
         "retention and NPS baseline set"],
        ["Q4 2027", "**Gate 3 — Repeatable engine**",
         "**500k cumulative ARR**; ~24 customers; playbook documented (pricing, install, support "
         "SLA); MOT 2028 readiness position taken; go/no-go for Vietnam"],
    ], [1.9, 4.0, 12.7], bold_cols=(0,), zebra=True)

    # ---------------------------------------------------------- 7. budget
    section(doc, "7", "BUDGET ENVELOPE (USD, ESTIMATE — REQUIRES COO / CFO VALIDATION)")
    table(doc, [
        ["Category", "2026 (Q3–Q4)", "2027", "Total"],
        ["Research, travel and local advisory", "15–25k", "15–25k", "30–50k"],
        ["Entity setup and compliance (Sdn Bhd, resident director, co-sec)",
         "10–18k", "6–10k", "16–28k"],
        ["Certification — SIRIM/MCMC type approval (2 SKUs) + MIROS TrackScore",
         "10–20k", "5–10k", "15–30k"],
        ["Headcount — country lead, 1 enterprise AE, 2 mid-market AEs, 1 CS/implementation",
         "20–30k", "110–155k", "130–185k"],
        ["Installer network setup and training", "—", "15–25k", "15–25k"],
        ["Marketing, events and channel enablement", "5–10k", "30–50k", "35–60k"],
        ["Hardware working capital (peak, ~3,000 units)", "—", "60–90k", "60–90k"],
        ["**Total cash envelope**", "**60–103k**", "**241–365k**", "**301–468k**"],
    ], [10.2, 2.8, 2.8, 2.8], align_right={1, 2, 3}, zebra=True)
    for c in range(4):
        shade(doc.tables[-1].cell(8, c), LIGHT)
    para(doc, "Approximately 0.77x burn-to-ARR at midpoint (~385k invested → 500k exit ARR). "
              "Product and engineering localisation is absorbed by the existing Indonesia team "
              "(allocation, not cash). Certification figures are the least reliable line here and "
              "are the first the Q3 2026 milestone must replace with quoted numbers.",
         size=SZ_NOTE, italic=True, color=C_GREY, before=2, after=0)

    # ---------------------------------------------------------- 8. risks
    section(doc, "8", "KEY RISKS & MITIGATION")
    table(doc, [
        ["Risk", "Severity", "Mitigation"],
        ["**Certification lead time** blocks legal shipment; first install slips past Q2 2027",
         "Critical",
         "File SIRIM/MCMC in Q4 2026; freeze SKU by November 2026; confirm lead time in Q3 2026"],
        ["**Sales capacity** — ~24 accounts across three selling quarters is not a one-AE plan",
         "High",
         "Country lead by Q4 2026; enterprise AE in Q1; two mid-market AEs by Q2/Q3 2027"],
        ["**Anchor concentration** — half the target sits in 3–4 logos; one loss is a "
         "~80k gap", "High",
         "Qualify 8+ anchors to close 4; monthly stage-gate pipeline review with sponsor"],
        ["**MOT pilot cohort closed** — five providers hold the reference position for a "
         "possible 2028 mandate", "High",
         "Test pilot access in Q3 2026; if closed, compete on TrackScore rating and APAD ICOP "
         "conformance"],
        ["**Installer capacity** — ~2,000 of ~3,000 units install in H2 2027 (~330/month, "
         "5–6 teams)", "Med-High",
         "Contract ≥4 partners by Q1 2027 with SLA capacity commitments, not goodwill"],
        ["**Entity and resident-director drag** — delays block local invoicing", "Medium",
         "Start Q4 2026 through a corporate services provider"],
    ], [7.6, 1.9, 9.1], zebra=True)

    # ---------------------------------------------------------- 9. assumptions
    section(doc, "9", "ASSUMPTIONS")
    bullets(doc, [
        "500k is **exit-2027 cumulative ARR**, not revenue recognised during 2027. All figures USD.",
        "USD 14/month blended ARPU is **conservative**: the market band implied by the MOT "
        "initiative is RM60–300 (~USD 13–65). Entering at the floor is a deliberate "
        "pricing choice; the mid-market half of the mix may support more.",
        "MIROS TrackScore is voluntary — the ≥4-star target is a commercial positioning "
        "decision, not a legal requirement.",
        "MOT telematics is **not yet mandatory**; the 2028 phase is conditional on industry "
        "readiness. The 500k target is not built on it.",
        "Localisation is delivered by the existing Indonesia product and engineering team; no "
        "Malaysian engineering hire is assumed.",
    ], size=SZ_TBL)

    # ---------------------------------------------------------- 10. governance
    section(doc, "10", "GOVERNANCE & APPROVAL")
    table(doc, [
        ["**Sponsor**", "Chief Operating Officer — owns the Gate 1 entry-mode decision and "
                        "the budget envelope"],
        ["**Project owner**", "Chief Strategy Officer — owns delivery, the risk register and "
                              "gate readiness"],
        ["**Workstreams**", "GTM & Sales [ ] · Product & Localisation [ ] · Hardware & "
                            "Certification [ ] · Legal & Entity [ ] · Finance [ ]"],
        ["**Cadence**", "Monthly steering with the sponsor; full C-Level review at each of the "
                        "three gates"],
    ], [3.0, 15.6], header=False)
    for r in range(4):
        shade(doc.tables[-1].cell(r, 0), LIGHT)

    para(doc, "", size=6, after=0)
    table(doc, [
        ["Approved by — Sponsor (COO)", "Date", "Approved by — Owner (CSO)", "Date"],
        ["", "", "", ""],
    ], [6.3, 3.0, 6.3, 3.0], size=SZ_TBL)
    for c in range(4):
        p = doc.tables[-1].cell(1, c).paragraphs[0]
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)

    doc.save(path)
    return path


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else "Malaysia_Expansion_Project_Charter_v1.docx"
    print(build(out))
